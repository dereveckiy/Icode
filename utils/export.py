import os
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook
from kivy.utils import platform

# Исходный заголовок (без нумерации)
HEADER = ["D", "L", "Qty", "Wood", "Sort", "V", "Price", "Total"]
# Новый заголовок с колонкой нумерации
NEW_HEADER = ["No."] + HEADER

def add_numbering(data):
    """
    Добавляет в каждую строку порядковый номер (начиная с 1) как первый элемент.
    """
    new_data = []
    for i, row in enumerate(data, start=1):
        new_data.append([str(i)] + list(row))
    return new_data

def compute_summary(filtered_data, header):
    """
    Вычисляет итоговые суммы для столбцов V и Total (если присутствуют).
    Возвращает итоговую строку, согласованную с header.
    """
    sum_v = 0.0
    sum_total = 0.0
    idx_v = None
    idx_total = None
    if "V" in header:
        idx_v = header.index("V")
    if "Total" in header:
        idx_total = header.index("Total")
    for row in filtered_data:
        if idx_v is not None:
            try:
                sum_v += float(row[idx_v])
            except (ValueError, IndexError):
                pass
        if idx_total is not None:
            try:
                sum_total += float(row[idx_total])
            except (ValueError, IndexError):
                pass
    # Для итоговой строки заполняем пустые ячейки, добавляем суммы в соответствующих колонках
    summary = []
    for col in header:
        if col == "V":
            summary.append(f"{sum_v:.6f}")
        elif col == "Total":
            summary.append(f"{sum_total:.2f}")
        else:
            summary.append("")
    # Можно в первый столбец поставить слово "Итого:" или оставить пустым
    if header:
        summary[0] = "Итого:"
    return summary

def filter_data(new_data, header_info, selected_columns):
    """
    Фильтрует заголовок и данные согласно выбранным опциям.
    Если выбрана опция "Метод", вставляет её перед остальными данными.
    Для остальных опций (Порода древесины, Сорт, Цена) удаляет соответствующие столбцы,
    если они не выбраны.
    """
    # Копируем базовый заголовок
    filtered_header = NEW_HEADER.copy()
    # Если пользователь выбрал экспорт столбца "Метод",
    # добавляем его после колонки "No." (индекс 1)
    if "Метод" in selected_columns:
        filtered_header.insert(1, "Method")
    # Соответствие опций и названий столбцов
    mapping = {
        "Порода древесины": "Wood",
        "Сорт": "Sort",
        "Цена": "Price"
    }
    # Удаляем столбцы, если опция не выбрана
    for option, col_name in mapping.items():
        if option not in selected_columns and col_name in filtered_header:
            filtered_header.remove(col_name)

    # Фильтрация данных
    filtered_data = []
    for row in new_data:
        new_row = row.copy()
        # Если "Метод" выбрана, вставляем её значение (берём из заголовочной информации)
        if "Метод" in selected_columns:
            new_row.insert(1, header_info.get("Метод", ""))
        # Формируем словарь исходного соответствия
        row_dict = dict(zip(NEW_HEADER, new_row[:len(NEW_HEADER)+1]))
        # Строим новую строку в порядке filtered_header
        filtered_row = [row_dict.get(col, "") for col in filtered_header]
        filtered_data.append(filtered_row)
    return filtered_header, filtered_data

def export_to_pdf(data, filename, header_info, selected_columns):
    """
    Экспортирует данные в PDF с учетом выбранных опций.
    """
    new_data = add_numbering(data)
    filtered_header, filtered_data = filter_data(new_data, header_info, selected_columns)
    summary = compute_summary(filtered_data, filtered_header)

    c = canvas.Canvas(filename, pagesize=A4)
    c.setFont("Helvetica", 10)
    width, height = A4
    x = 40
    y = height - 40
    line_height = 15

    # Добавляем заголовок документа из header_info (если есть)
    header_lines = []
    for key in ["Номер накладной", "Поставщик", "Контрагент"]:
        if header_info.get(key):
            header_lines.append(f"{key}: {header_info[key]}")
    for hl in header_lines:
        c.drawString(x, y, hl)
        y -= line_height
    y -= line_height

    header_line = "    ".join(filtered_header)
    c.drawString(x, y, header_line)
    y -= line_height * 1.5

    for row in filtered_data:
        row_line = "    ".join(str(item) for item in row)
        c.drawString(x, y, row_line)
        y -= line_height
        if y < 40:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - 40

    summary_line = "    ".join(str(item) for item in summary)
    c.drawString(x, y - line_height, summary_line)
    c.save()

def export_to_word(data, filename, header_info, selected_columns):
    """
    Экспортирует данные в Word документ (.docx) с учетом выбранных опций.
    """
    new_data = add_numbering(data)
    filtered_header, filtered_data = filter_data(new_data, header_info, selected_columns)
    summary = compute_summary(filtered_data, filtered_header)

    document = Document()
    # Добавляем заголовочную информацию
    for key in ["Номер накладной", "Поставщик", "Контрагент"]:
        if header_info.get(key):
            document.add_paragraph(f"{key}: {header_info[key]}")
    table = document.add_table(rows=1, cols=len(filtered_header))
    hdr_cells = table.rows[0].cells
    for j, col_name in enumerate(filtered_header):
        hdr_cells[j].text = str(col_name)
    for row in filtered_data:
        cells = table.add_row().cells
        for j in range(len(filtered_header)):
            cells[j].text = str(row[j])
    sum_row = table.add_row().cells
    for j in range(len(filtered_header)):
        sum_row[j].text = summary[j] if j < len(summary) else ""
    document.save(filename)

def export_to_excel(data, filename, header_info, selected_columns):
    """
    Экспортирует данные в Excel файл (.xlsx) с учетом выбранных опций.
    """
    new_data = add_numbering(data)
    filtered_header, filtered_data = filter_data(new_data, header_info, selected_columns)
    summary = compute_summary(filtered_data, filtered_header)
    wb = Workbook()
    ws = wb.active

    # Записываем заголовочную информацию в шапку документа
    row_num = 1
    for key in ["Номер накладной", "Поставщик", "Контрагент"]:
        if header_info.get(key):
            ws.cell(row=row_num, column=1, value=f"{key}: {header_info[key]}")
            row_num += 1
    row_num += 1

    ws.append(filtered_header)
    for row in filtered_data:
        ws.append(row)
    ws.append(summary)
    wb.save(filename)

def share_file(filename):
    """
    Функция для "поделиться" файлом.
    
    Для Android используется plyer.share, для Windows открывается проводник с выделенным файлом.
    """
    full_path = os.path.abspath(filename)
    if platform == "android":
        try:
            from plyer import share
            share.share(title="Поделиться", text="Смотрите файл", file_path=full_path)
        except Exception as e:
            print("Ошибка при попытке поделиться на Android:", e)
    elif platform == "win":
        try:
            import subprocess
            subprocess.Popen(["explorer", "/select,", full_path])
        except Exception as e:
            print("Ошибка при попытке поделиться на Windows:", e)
    else:
        print("Платформа не поддерживается для функции 'Поделиться'.")